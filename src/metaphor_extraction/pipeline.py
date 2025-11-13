"""
Metaphor Extraction Pipeline
Steps followed:
    1. Load configuration
    2. Setup text splitter
    3. Setup metaphor extraction chain
    4. Read .docx files
    5. Split documents into chunks
    6. Run metaphor extraction and validation
"""

import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from .chain import create_metaphor_chain, get_model
from .structured_output import MetaphorsResponse
from .validation import SimpleQuoteValidator

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],
)
logger = logging.getLogger(__name__)


class MetaphorPipeline:
    def __init__(self, config_path: Optional[str] = None):
        """Initialize the pipeline with configuration"""
        if config_path is None:
            config_path = str(Path(__file__).parent / "config.yaml")

        self.config = self.load_config(config_path)
        self.provider = self.config["default_provider"]

        self.setup_text_splitter()
        self.setup_metaphor_chain()

    def load_config(self, config_path: str) -> Dict:
        """Load configuration from YAML file"""

        if os.path.exists(config_path):
            with open(config_path, "r") as f:
                config = yaml.safe_load(f)
        else:
            raise FileNotFoundError(
                f"Configuration file not found: {config_path}"
            )
        return config

    def setup_text_splitter(self):
        """Setup text splitter"""
        chars_per_token = self.config["text_splitter"].get(
            "chars_per_token", 4
        )

        chunk_size = (
            self.config["text_splitter"]["max_tokens"] * chars_per_token
        )
        chunk_overlap = (
            self.config["text_splitter"]["overlap_tokens"] * chars_per_token
        )

        # Regex separators for H-number structure
        # Order matters: most specific to least specific
        separators = [
            # Match main headers like **H0**, **H1**, **H2**, etc.
            r"\n\*\*H\d+\*\*",
            # Match subsections like H0-1, H1-2, H2-10, etc.
            r"\nH\d+-\d+",
            # Match speaker turns (with colon)
            r"\n[PON]:\s",  # P: (Patient), O: (Onderzoeker), N: (Naaste)
            # Fallback separators
            r"\n\n",  # Double newline
            r"\n",  # Single newline
            r"\.\s",  # Sentence boundary
            r"\s",  # Word boundary
            "",  # Character boundary (last resort)
        ]

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            length_function=len,
            keep_separator="start",
            is_separator_regex=True,
        )

    def setup_metaphor_chain(self):
        """Setup the metaphor extraction chain"""
        model_name = self.config["providers"][self.provider]["model_name"]
        base_url = self.config["providers"][self.provider].get(
            "base_url", None
        )
        options = self.config["providers"][self.provider].get("options", {})

        output_format = MetaphorsResponse.model_json_schema()

        model = get_model(
            self.provider,
            base_url=base_url,
            model_name=model_name,
            format=output_format,
            options=options,
        )

        self.metaphor_chain = create_metaphor_chain(
            model=model, prompt_id=self.config.get("prompt_id", 1)
        )
        self.quote_validator = SimpleQuoteValidator()
        logger.info("Metaphor extraction chain setup complete.")

    def read_docx(self, file_path: str) -> str:
        """Read text from a .docx file"""
        doc = DocxDocument(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)

    def split_document_for_llm(self, text):
        """
        Split document into chunks suitable for Llama3.1
        Llama3.1 context window is typically 8K-32K tokens
        Leave room for prompt + response
        """
        # Create Document objects
        docs = [Document(page_content=text)]
        chunks = self.text_splitter.split_documents(docs)

        return chunks

    def run(self, file_path: str) -> List[Dict[str, Any]]:
        """Run the metaphor extraction pipeline on a given text file"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Input file not found: {file_path}")

            text = self.read_docx(file_path)

            # Split document into chunks
            chunks = self.split_document_for_llm(text)
            logging.info(f"Document split into {len(chunks)} chunks.")

            metaphors = []

            for idx, chunk in enumerate(chunks):
                logger.info(
                    f"Running metaphor extraction chain on chunk-{idx}..."
                )
                response: MetaphorsResponse = self.metaphor_chain.invoke(
                    {"document_text": chunk.page_content}
                )

                logging.info(
                    f"Chunk-{idx} processed. Extracted {len(response.metaphors)} metaphors."
                )

                if len(response.metaphors) > 0:
                    # Validate quotes in the response
                    for metaphor in response.metaphors:
                        logging.info(
                            f"Validating quote: {metaphor.original_quote}"
                        )
                        validation_scores = (
                            self.quote_validator.validate_quote(
                                quote=metaphor.original_quote,
                                source_text=chunk.page_content,
                            )
                        )
                        metaphor_response = metaphor.model_dump()
                        metaphor_response["exact_match"] = validation_scores
                        metaphors.append(metaphor_response)

            logger.info(
                f"Metaphor extraction and validation complete. Total {len(metaphors)} metaphors found."
            )

            return metaphors

        except Exception as e:
            logger.error(f"Pipeline failed: {str(e)}")
            raise e
